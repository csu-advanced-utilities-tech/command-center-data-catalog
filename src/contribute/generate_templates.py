import pandas as pd
from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).resolve().parents[2]

VALIDATED_DIR = BASE_DIR / "output" / "validated"
META_DIR = BASE_DIR / "metadata"
TEMPLATE_DIR = BASE_DIR / "templates"

TEMPLATE_DIR.mkdir(exist_ok=True)

def generate_table_template():
    tech = pd.read_csv(VALIDATED_DIR / "technical_columns_validated.csv")

    # Drive coverage from technical reality
    tables = (
        tech[["OWNER", "TABLE_NAME"]]
        .drop_duplicates()
        .rename(columns={"OWNER": "owner", "TABLE_NAME": "table_name"})
        .sort_values(["owner", "table_name"])
        .reset_index(drop=True)
    )

    # Load existing business metadata if present
    business = pd.read_csv(META_DIR / "table_metadata.csv")

    merged = tables.merge(
        business,
        how="left",
        on=["owner", "table_name"]
    )

    # Keep ONLY what appears on the UI
    ui_columns = [
        "owner",
        "table_name",
        "business_name",
        "description",
        "grain",
        "primary_key",
        "notes",
    ]

    merged = merged[ui_columns]

    out = TEMPLATE_DIR / "table_metadata_TEMPLATE.xlsx"
    merged.to_excel(out, index=False)
    print(f"✅ Generated {out.name}")

def generate_column_template():
    tech = pd.read_csv(VALIDATED_DIR / "technical_columns_validated.csv")

    columns = (
        tech[["OWNER", "TABLE_NAME", "COLUMN_NAME"]]
        .rename(
            columns={
                "OWNER": "owner",
                "TABLE_NAME": "table_name",
                "COLUMN_NAME": "column_name",
            }
        )
        .sort_values(["owner", "table_name", "column_name"])
        .reset_index(drop=True)
    )

    business = pd.read_csv(META_DIR / "column_metadata.csv")

    merged = columns.merge(
        business,
        how="left",
        on=["owner", "table_name", "column_name"]
    )

    # Keep ONLY what appears on the UI
    ui_columns = [
        "owner",
        "table_name",
        "column_name",
        "business_name",
        "description",
        "example_value",
        "is_sensitive",
        "notes",
    ]

    merged = merged[ui_columns]

    out = TEMPLATE_DIR / "column_metadata_TEMPLATE.xlsx"
    merged.to_excel(out, index=False)
    print(f"✅ Generated {out.name}")

def generate_readme():
    doc = Document()
    doc.add_heading("Data Catalog Contribution Guide", level=1)

    doc.add_paragraph(
        "These Excel templates mirror the Data Catalog website exactly. "
        "Each row corresponds to what you see on the site."
    )

    doc.add_heading("How to Edit", level=2)
    doc.add_paragraph("• Fill in blank cells where metadata is missing")
    doc.add_paragraph("• You may update existing business descriptions")
    doc.add_paragraph("• Do NOT rename tables or columns")
    doc.add_paragraph("• Do NOT delete rows")

    doc.add_heading("Submission Process", level=2)
    doc.add_paragraph(
        "When finished, upload the edited files back to SharePoint. "
        "Changes will be reviewed and published by the catalog owner."
    )

    out = TEMPLATE_DIR / "README_HOW_TO_EDIT.docx"
    doc.save(out)
    print("✅ Generated README_HOW_TO_EDIT.docx")

def main():
    generate_table_template()
    generate_column_template()
    generate_readme()

if __name__ == "__main__":
    main()
